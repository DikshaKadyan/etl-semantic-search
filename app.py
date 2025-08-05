import streamlit as st
import fitz  # PDF
import docx
import pandas as pd
import json
import numpy as np
from langchain.text_splitter import RecursiveCharacterTextSplitter
from sentence_transformers import SentenceTransformer
from langchain_community.vectorstores import InMemoryVectorStore
from langchain.schema import Document

# ---------- Initialize ----------
embedder = SentenceTransformer("all-MiniLM-L6-v2")
splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)
vector_store = None


# ---------- Extract ----------
def extract_text(file):
    if file.name.endswith(".pdf"):
        doc = fitz.open(stream=file.read(), filetype="pdf")
        return " ".join([page.get_text() for page in doc])
    elif file.name.endswith(".docx"):
        doc = docx.Document(file)
        return " ".join([para.text for para in doc.paragraphs])
    elif file.name.endswith(".txt"):
        return file.read().decode("utf-8")
    elif file.name.endswith(".csv"):
        df = pd.read_csv(file)
        return df.to_string()
    elif file.name.endswith(".json"):
        data = json.load(file)
        return json.dumps(data, indent=2)
    else:
        return ""


# ---------- Transform ----------
def transform_text(text):
    chunks = splitter.split_text(text)
    embeddings = embedder.encode(chunks, show_progress_bar=False)
    return list(zip(chunks, embeddings))


# ---------- Load into In-Memory Vector Store ----------
def load_data(chunks_with_embeddings):
    global vector_store
    docs = [Document(page_content=chunk) for chunk, _ in chunks_with_embeddings]
    vectors = [emb for _, emb in chunks_with_embeddings]
    vector_store = InMemoryVectorStore.from_embeddings(docs, vectors)


# ---------- Query ----------
def query_data(query, top_k=3):
    if not vector_store:
        return []
    query_vector = embedder.encode([query])[0]
    results = vector_store.similarity_search_by_vector(query_vector, k=top_k)
    return [doc.page_content for doc in results]


# ---------- Streamlit UI ----------
st.title("📑 ETL Pipeline: Semantic Search in Multi-format Documents")
st.write("Upload a file (PDF, DOCX, TXT, CSV, JSON) and ask questions about it.")

uploaded_file = st.file_uploader("Upload a document", type=["pdf", "docx", "txt", "csv", "json"])

if uploaded_file:
    text = extract_text(uploaded_file)
    st.success("✅ File uploaded and text extracted")

    chunks_with_embeddings = transform_text(text)
    load_data(chunks_with_embeddings)
    st.success("✅ Text processed and stored in in-memory vector store")

    query = st.text_input("Ask a question about the document:")
    if query and vector_store:
        results = query_data(query, top_k=3)
        st.write("### 🔎 Top Results")
        for doc in results:
            st.write("-", doc)
